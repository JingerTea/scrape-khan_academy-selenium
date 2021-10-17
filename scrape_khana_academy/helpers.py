import time
import os
import re
import shutil
from webdriver_manager.chrome import ChromeDriverManager # webdriver_manager
from selenium import webdriver # selenium
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from docx import Document # python-docx
from docx.shared import Inches # python-docx
import scrape_khana_academy.constants as c

def create_folder(path, name, notify=True):
    if not os.path.exists(f'{path}\\{name}'):
        os.makedirs(f'{path}\\{name}')
        if notify:
            print(f"Folder created: {name}")
    else:
        if notify:
            print(f"Folder existed: {name}")


def khan_academy_scraper(driver, link, path=os.getcwd()):
    # Initialize
    create_folder(path, c.MAIN_FOLDER)
    create_folder(path, c.TEMP_FOLDER, notify=False)
    # Open Khan Academy Website
    driver.get(link)
    grade_title = driver.find_element_by_xpath(c.GRADE).text
    grade_title = re.sub('[^A-Za-z0-9]+', '_', grade_title)
    grade_title = f'{c.MAIN_FOLDER}\\{grade_title}'
    create_folder(path, grade_title)

    # Loop Unit
    for i in range(len(driver.find_elements_by_xpath(c.UNIT))):
        # Create Unit Folder
        unit_folder = driver.find_element_by_xpath(f'({c.UNIT})[{i+1}]').text
        # Remove Forbidden Printable ASCII Characters
        unit_folder = re.sub('[^A-Za-z0-9]+', '_', f"{i+1}_{unit_folder}")
        unit_folder = f'{grade_title}\\{unit_folder}'
        create_folder(path, unit_folder)

        driver.find_element_by_xpath(f'({c.UNIT})[{i+1}]').click()
        time.sleep(3)
        unit_url = driver.current_url

        # Loop practice
        for i in range(len(driver.find_elements_by_xpath(c.PRACTICE_BUTTON))):
            practice_title = driver.find_element_by_xpath(f'({c.PRACTICE_TITLE})[{i+1}]').text
            practice_filename = re.sub('[^A-Za-z0-9]+', '_', f"{i+1}_{practice_title}")

            if os.path.exists(f'{unit_folder}\\{practice_filename}.docx'):
                print(f"Docx existed: {practice_filename}.docx")
            else:
                # Practice - Start
                driver.find_element_by_xpath(f'({c.PRACTICE_BUTTON})[{i+1}]').click()
                time.sleep(2)
                # Practice - Let's Go
                driver.find_element_by_xpath(f'({c.PRACTICE_DIALOG_CHECK})[last()]').click()

                # Loop Question
                for i in range(len(driver.find_elements_by_xpath(c.PRACTICE_PROGRESS))):
                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, c.PRACTICE_PROBLEM))
                    )
                    # Screenshot Problem
                    image = driver.find_element_by_xpath(c.PRACTICE_PROBLEM)
                    image.screenshot(f'{c.TEMP_FOLDER}\\Question{i+1}.png')
                    # Mutiple-Choice Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_CHOICE)) > 0:
                        # print("Mutiple-Choice")
                        driver.find_element_by_xpath(c.PRACTICE_CHOICE).click()
                    # Short Answer Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_TEXTBOX)) > 0:
                        # print("Short Answer")
                        for i in range(len(driver.find_elements_by_xpath(c.PRACTICE_TEXTBOX))):
                            driver.find_element_by_xpath(f'({c.PRACTICE_TEXTBOX})[{i + 1}]').send_keys("0")
                    # Dropdown Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_DROPDOWN)) > 0:
                        # print("Dropdown")
                        for i in range(len(driver.find_elements_by_xpath(c.PRACTICE_DROPDOWN))):
                            driver.find_element_by_xpath(f'({c.PRACTICE_DROPDOWN})[{i + 1}]').click()
                            time.sleep(1)
                            driver.find_element_by_xpath(f'({c.PRACTICE_DROPDOWN_CHOICE})[last()]').click()
                    #  Equation Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_EQUATION)) > 0:
                        # print("Equation")
                        for i in range(len(driver.find_elements_by_xpath(c.PRACTICE_EQUATION))):
                            driver.find_element_by_xpath(f'({c.PRACTICE_EQUATION})[{i + 1}]').click()
                            # Input without target element
                            action = ActionChains(driver)
                            action.send_keys('0').perform()
                    # Table Selection Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_TABLE)) > 0:
                        # print("Table Selection")
                        driver.find_element_by_xpath(c.PRACTICE_TABLE).click()
                    # Graph Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_PLOT)) > 0:
                        # print("Graph")
                        source = driver.find_element_by_xpath(c.PRACTICE_PLOT)
                        target = driver.find_element_by_xpath(c.PRACTICE_GRAPH)
                        action = ActionChains(driver) # Reset ActionChains required or else element not found
                        action.click_and_hold(source)
                        action.move_to_element_with_offset(target, 50, 50)
                        action.click().perform()
                    # Sort Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_SORT)) > 0:
                        # print("Sort")
                        driver.find_element_by_xpath(c.PRACTICE_SORT).click()
                        driver.find_element_by_xpath(c.PRACTICE_SORT).click()
                    # Point Problem
                    if len(driver.find_elements_by_xpath(c.PRACTICE_POINT)) > 0:
                        # print("Point")
                        for i in range(len(driver.find_elements_by_xpath(c.PRACTICE_POINT))):
                            driver.find_element_by_xpath(f'{c.PRACTICE_POINT}[{i+1}]').click()
                            time.sleep(1)
                            driver.find_element_by_xpath(c.PRACTICE_POINT_CHOICE).click()


                    driver.find_element_by_xpath(f'({c.PRACTICE_DIALOG_CHECK})[last()]').click()
                    time.sleep(2)

                    # Wrong Answer
                    if len(driver.find_elements_by_xpath(c.PRACTICE_MOVEON)) > 0:
                        driver.find_element_by_xpath(c.PRACTICE_MOVEON).click()
                    else:
                        driver.find_element_by_xpath(f'({c.PRACTICE_DIALOG_CHECK})[last()-1]').click()
                    time.sleep(2)

                driver.get(unit_url)
                time.sleep(3)

                # Create Docx
                picture_names = os.listdir(c.TEMP_FOLDER)
                document = Document()
                sections = document.sections
                for section in sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

                document.add_heading({practice_title}, 0)
                for picture_name in picture_names:
                    file_path = f'{c.TEMP_FOLDER}\\{picture_name}'
                    document.add_picture(file_path, width=Inches(6.1))
                    document.add_paragraph("____________________________________________________"
                                           "_____________________________________________________")
                    document.add_paragraph("Scratch paper:")
                    if picture_name == picture_names[-1]:
                        pass
                    else:
                        document.add_page_break()
                document.save(f'{unit_folder}\\{practice_filename}.docx')
                print(f"Docx created: {practice_filename}.docx")
                # Clean up
                shutil.rmtree(c.TEMP_FOLDER)
                create_folder(path, c.TEMP_FOLDER, notify=False)
        driver.get(link)
        time.sleep(3)